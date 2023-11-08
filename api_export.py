# api_export.py
import os
import uuid
import string
# import re

from fastapi import HTTPException
from pydantic import BaseModel
from typing import List
from docx import Document
from docx.shared import Cm

import dbconnect
from questionbank_launch import app
import commonfuncs as cf
from api_files import image_size

root = os.path.dirname(__file__)
ANSWER_TYPES = ("MCQ_single", "InQuestion", "TrueFalse", "MTF")
UPLOAD_FOLDER = os.environ.get("UPLOAD_FOLDER", "data/uploads")
uploadFolder = os.path.join(root, UPLOAD_FOLDER)

EXPORT_FOLDER = os.environ.get("EXPORT_FOLDER","data/exports")
exportFolder = os.path.join(root, EXPORT_FOLDER)
os.makedirs(exportFolder, exist_ok=True)

alphabet_list = [letter for letter in string.ascii_lowercase]

PAGE_W = 21.0
PAGE_H = 29.7    
DPC = 37.8 # dots per cm, equivalent to 96 dpi

#####################
# FUNCTIONS

def insert_image(filename, p, embeds):
    global uploadFolder, DPC, PAGE_H, PAGE_W
    image_path = os.path.join(uploadFolder, embeds.get(filename))
    w, h = image_size(image_path)
    
    w_cm = round(w / DPC,2)
    h_cm = round(h / DPC,2)
    cf.logmessage("image width, height:",w_cm, h_cm)

    if w_cm > PAGE_W*0.65:
        w_cm = PAGE_W*0.65
        cf.logmessage("Image too wide! New width:",w_cm)
        p.add_run().add_picture(image_path, width=Cm(w_cm))

    elif  h_cm > PAGE_H*0.65:
        h_cm = PAGE_H*0.65
        cf.logmessage("Image too tall! New height:",h_cm)
        p.add_run().add_picture(image_path, height=Cm(h_cm))
    else:
        p.add_run().add_picture(image_path, width=Cm(w_cm))
    return


def paraHandleImages(s, doc, style=None, embeds={}):
    global uploadFolder
    img_placeholders, non_placeholder_parts = cf.split_string_with_placeholders(s)

    # only string, no images
    if len(img_placeholders) == 0:
        p = doc.add_paragraph(s, style=style)
        return p
    
    # if only images
    if len(non_placeholder_parts) == 0:
        p = doc.add_paragraph("", style=style)
        for N, filename in enumerate(img_placeholders):
            insert_image(filename, p, embeds)
        return p
    
    # if sentence surrounging image(s)
    if len(non_placeholder_parts) > len(img_placeholders):
        p = doc.add_paragraph("", style=style)
        for N, filename in enumerate(img_placeholders):
            p.add_run(non_placeholder_parts[N])
            insert_image(filename, p, embeds)
        p.add_run(non_placeholder_parts[-1])
        return p
    
    elif len(img_placeholders) == len(non_placeholder_parts):
        if s.startswith(r'{{'):
            p = doc.add_paragraph("", style=style)
            for N, filename in enumerate(img_placeholders):
                insert_image(filename, p, embeds)
                p.add_run(non_placeholder_parts[N])
            return p
        else:
            p = doc.add_paragraph("", style=style)
            for N, filename in enumerate(img_placeholders):
                p.add_run(non_placeholder_parts[N])
                insert_image(filename, p, embeds)
            return p

    else:
        # images surrounding sentences
        p = doc.add_paragraph("", style=style)
        for N, txt in enumerate(non_placeholder_parts):
            insert_image(img_placeholders[N], p, embeds)
            p.add_run(txt)
        image_path = os.path.join(uploadFolder, embeds.get(img_placeholders[-1]))
        p.add_run().add_picture(image_path)
        insert_image(img_placeholders[-1], p, embeds)
        return p

####

def writeTrueFalse(q, doc, embeds={}):
    for N, statement in enumerate(q['statements']):
        p = paraHandleImages(f"{N+1}) {statement} - ______", doc, embeds=embeds)
        p.paragraph_format.left_indent = Cm(0.4)
    doc.add_paragraph("")

def writeMCQ_single(q, doc, embeds={}):
    for N, choice in enumerate(q['choices']):
        p = paraHandleImages(choice, doc, style='List Bullet', embeds=embeds)
    
    doc.add_paragraph("")

def writeMTF(q, doc, embeds={}):
    table = doc.add_table(rows=1, cols=2)

    # Set the column widths equally
    for cell in table.columns[0].cells:
        cell.width = Cm(21.0 * 0.4)

    row = table.rows[0]
    for N, (val1, val2) in enumerate(zip(q['left_side'], q['right_side'])):
        cell1 = paraHandleImages(f"{alphabet_list[N]}) {val1}", row.cells[0])
        cell2 = paraHandleImages(f"{N+1}) {val2}", row.cells[1])
    
    doc.add_paragraph("")


#####################
# API CALLS

class export_payload(BaseModel):
    question_ids: List[int]

@app.post("/api/questions/export", tags=["questions"])
def export(r: export_payload):
    global exportFolder, alphabet_list, PAGE_W, PAGE_H
    
    idsList = ",".join([str(x) for x in r.question_ids])
    s1 = f"""select t1.*,
    t2.subject_name, t2.topic_name, t2.subtopic_name
    from questionbank as t1 
    left join topics as t2
    on t1.subtopic_id = t2.subtopic_id
    where t1.id in ({idsList})
    """
    df1 = dbconnect.makeQuery(s1, output="df")
    
    if not len(df1):
        return {"message":"no valid selection"}
    
    # https://python-docx.readthedocs.io/en/latest/index.html
    doc = Document()

    # Set page size
    section = doc.sections[0]
    section.page_width = Cm(PAGE_W)  
    section.page_height = Cm(PAGE_H)  
    doc.add_heading('Question Bank', 0)
    
    for N,r in df1.iterrows():
        q = cf.yaml2dict(r['content'])
        embeds = cf.embeds2dict(q.get('embeds',[]))

        p = paraHandleImages(f"{N+1}. {q['question']}", doc, embeds=embeds)

        if q['answer_type'] == "TrueFalse":
            writeTrueFalse(q, doc, embeds=embeds)
        elif q['answer_type'] == "MCQ_single":
            writeMCQ_single(q, doc, embeds=embeds)
        elif q['answer_type'] == "MTF":
            writeMTF(q, doc, embeds=embeds)
        else:
            # default: assume InQuestion
            doc.add_paragraph("")

    filename = f"questionpaper_{str(uuid.uuid4())}.docx"
    doc.save(os.path.join(exportFolder, filename))

    return {"filename": filename}

